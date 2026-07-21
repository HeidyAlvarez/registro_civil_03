"""Genera el documento Word con las tarjetas CRC corregidas."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


TARJETAS = [
    {
        "numero": 15,
        "titulo": "Usuario",
        "nombre": "UsuarioRegistroCivil",
        "tipo": "Modelo — autenticación",
        "capa": "Capa: Datos",
        "responsabilidades": [
            "Heredar los campos base de AbstractUser de Django (usuario, contraseña, email, etc.)",
            "Agregar el campo rol con valores almacenados: ADMIN, OFICIAL o CAPTURISTA (default: CAPTURISTA)",
            "Exponer la etiqueta legible del rol mediante get_rol_display()",
            "Sincronizar el rol con los grupos de Django al guardar el usuario (Administrador, oficial, Capturista)",
            "Reemplazar al usuario estándar mediante AUTH_USER_MODEL en settings.py",
        ],
        "colaboradores": [
            "→ Abstracción de Usuario de Django (lo hereda)",
            "→ Cita (usuario_atendio referencia este modelo)",
            "→ PagoCaja (cajero referencia este modelo)",
            "→ BitacoraAuditoria (usuario referencia este modelo)",
            "→ CorteCajaDiario (cerrado_por referencia este modelo)",
        ],
        "nota": (
            "Esta tabla describe el modelo UsuarioRegistroCivil de la capa de datos. Extiende el usuario "
            "de Django con el campo operativo rol. En la práctica, la autorización del sistema se resuelve "
            "principalmente por grupos de Django, sincronizados desde ese rol al crear o editar usuarios."
        ),
    },
    {
        "numero": 16,
        "titulo": "Login",
        "nombre": "Autenticación y sesión",
        "tipo": "Lógica de acceso — autenticación",
        "capa": "Capa: Lógica de Negocio",
        "responsabilidades": [
            "Autenticar al usuario verificando usuario y contraseña mediante Django Auth",
            "Validar política de contraseña: mínimo 8 caracteres, una mayúscula y un número",
            "Iniciar y cerrar la sesión del usuario",
            "Redirigir al panel correcto según rol/grupo tras iniciar sesión",
            "Registrar en bitácora cada acceso exitoso y cada intento fallido con IP, usuario, fecha y hora",
            "Cerrar sesión por inactividad según SESSION_COOKIE_AGE",
            "Proteger vistas internas rechazando peticiones sin sesión activa",
        ],
        "colaboradores": [
            "→ UsuarioRegistroCivil (verifica credenciales y obtiene el rol)",
            "→ BitacoraAuditoria (registra accesos e intentos fallidos)",
            "→ Grupos de Django (determinan permisos operativos)",
        ],
        "nota": (
            "Implementada en autenticacion/servicios/login.py (clase Login). Django Auth sigue "
            "manejando credenciales y sesión; Login centraliza redirección por rol y mensajes de "
            "auditoría. El login del admin delega en esta clase."
        ),
    },
    {
        "numero": 17,
        "titulo": "Sección trámite",
        "nombre": "SeccionTramite",
        "tipo": "Modelo — citas",
        "capa": "Capa: Datos",
        "responsabilidades": [
            "Almacenar el nombre de la sección o categoría (VARCHAR 200, UNIQUE)",
            "Agrupar trámites por tipo para facilitar la navegación del ciudadano",
        ],
        "colaboradores": [
            "→ Tramite (una sección contiene muchos trámites)",
        ],
        "nota": (
            "Esta tabla describe únicamente el modelo SeccionTramite, encargado de organizar el catálogo "
            "de trámites por categorías."
        ),
    },
    {
        "numero": 18,
        "titulo": "Trámite",
        "nombre": "Tramite",
        "tipo": "Modelo — citas",
        "capa": "Capa: Datos",
        "responsabilidades": [
            "Almacenar nombre, duración en minutos, costo, documentos requeridos y estado activo/inactivo",
            "Relacionarse con una SeccionTramite mediante llave foránea",
            "Conservar la duración en minutos (duracion_minutos) como dato base del catálogo",
            "Impedir eliminación o desactivación cuando existan citas futuras que lo comprometan",
        ],
        "colaboradores": [
            "→ SeccionTramite (pertenece a una sección)",
            "→ Cita (cada cita referencia un trámite)",
        ],
        "nota": (
            "Esta tabla describe el modelo Tramite. La duración legible para interfaz (por ejemplo, "
            "1h 30min) se obtiene en la capa de presentación/API a partir de duracion_minutos; no existe "
            "un método persistente en el modelo."
        ),
    },
    {
        "numero": 19,
        "titulo": "Cita",
        "nombre": "Cita",
        "tipo": "Modelo central — citas",
        "capa": "Capa: Datos",
        "responsabilidades": [
            "Almacenar CURP (CHAR 18), nombre del ciudadano (VARCHAR 150) y código postal (CHAR 5)",
            "Almacenar dirección completa y datos adicionales del trámite en JSONField cuando aplique",
            "Guardar fecha, hora y estado actual (VARCHAR 15): PENDIENTE, ASISTIDA, PAGADA, FINALIZADA, CANCELADA",
            "Almacenar código QR único (VARCHAR 64, UNIQUE) e imagen del QR (ImageField)",
            "Registrar qué usuario atendió la cita (FK nullable a UsuarioRegistroCivil)",
            "Guardar la fecha y hora de creación del registro (creado_el, DATETIME)",
            "Forzar que la CURP siempre se almacene en mayúsculas",
            "Validar empalmes de horario considerando la duración del trámite",
        ],
        "colaboradores": [
            "→ Tramite (cada cita pertenece a un trámite)",
            "→ UsuarioRegistroCivil (registra quién atendió)",
            "→ PagoCaja (relación 1:1)",
            "→ BitacoraAuditoria (sus cambios pueden auditarse desde la lógica de negocio)",
        ],
        "nota": (
            "Esta tabla presenta el modelo central Cita, que concentra la información del ciudadano, "
            "la agenda, el estado operativo, el QR y la trazabilidad básica del trámite."
        ),
    },
    {
        "numero": 20,
        "titulo": "Pago caja",
        "nombre": "PagoCaja",
        "tipo": "Modelo — citas",
        "capa": "Capa: Datos",
        "responsabilidades": [
            "Almacenar monto cobrado (DECIMAL 10,2), fecha y hora del pago (DATETIME)",
            "Relacionarse de forma 1:1 con una Cita en estado ASISTIDA",
            "Registrar al cajero que realizó el cobro (FK a UsuarioRegistroCivil)",
            "Indicar mediante campo booleano si el pago ya fue incluido en el corte del día (corte_cierre_listo)",
            "Impedir alteración o eliminación cuando el pago ya pertenece a un corte cerrado",
        ],
        "colaboradores": [
            "→ Cita (relación 1:1; cada pago pertenece a una sola cita)",
            "→ UsuarioRegistroCivil (registra al cajero responsable)",
            "→ CorteCajaDiario (agrupa pagos en el cierre diario)",
        ],
        "nota": (
            "Esta tabla describe el modelo PagoCaja, encargado de registrar cada cobro en ventanilla "
            "y su inclusión posterior en el corte diario."
        ),
    },
    {
        "numero": 21,
        "titulo": "Bitácora auditoría",
        "nombre": "BitacoraAuditoria",
        "tipo": "Modelo de auditoría — citas",
        "capa": "Capa: Datos",
        "responsabilidades": [
            "Almacenar acción (VARCHAR 25), descripción (TEXT), fecha y hora (DATETIME) e IP (GenericIPAddressField)",
            "Relacionarse con el usuario responsable (FK nullable)",
            "Ser completamente inmutable: bloquear cualquier intento de modificación o borrado",
        ],
        "colaboradores": [
            "→ UsuarioRegistroCivil (referencia quién realizó la acción)",
        ],
        "nota": (
            "Esta tabla presenta el modelo BitacoraAuditoria, encargado de conservar de forma inmutable "
            "las acciones críticas del sistema. El campo accion utiliza un catálogo controlado "
            "(LOGIN, TRANSACCION, MODIFICACION_CITA, etc.)."
        ),
    },
    {
        "numero": 22,
        "titulo": "Cita",
        "nombre": "Gestión de citas",
        "tipo": "Lógica principal — citas",
        "capa": "Capa: Lógica de Negocio",
        "responsabilidades": [
            "Validar que la CURP tenga exactamente 18 caracteres y cumpla el formato oficial",
            "Validar que el código postal tenga exactamente 5 dígitos numéricos",
            "Verificar disponibilidad del horario antes de crear la cita",
            "Crear la cita y generar su QR al confirmar el agendamiento",
            "Gestionar datos adicionales del trámite cuando aplique (por ejemplo, registro de nacimiento)",
            "Listar citas del día ordenadas por hora para paneles operativos",
            "Gestionar el flujo de estados: PENDIENTE → ASISTIDA → PAGADA → FINALIZADA",
            "Permitir cancelaciones solo a usuarios con rol Oficial o Administrador",
        ],
        "colaboradores": [
            "→ Cita (crea y consulta registros)",
            "→ Tramite (obtiene duración y costo)",
            "→ Calendario / utilidades de horario",
            "→ QR (generación e imagen)",
            "→ Bitacora (registra acciones del personal autorizado)",
            "→ UsuarioRegistroCivil (verifica rol para cancelaciones)",
        ],
        "nota": (
            "Implementada en citas/servicios/cita.py (clase Cita, importada como CitaNegocio). "
            "Las vistas HTTP en citas/views.py delegan en esta clase. Las citas del portal no "
            "generan bitácora porque no hay usuario autenticado en ese flujo."
        ),
    },
    {
        "numero": 23,
        "titulo": "Calendario",
        "nombre": "Control de horarios",
        "tipo": "Lógica de negocio — citas",
        "capa": "Capa: Lógica de Negocio",
        "responsabilidades": [
            "Verificar que un horario esté libre considerando la duración del trámite",
            "Detectar empalmes mediante validación de intervalos en la lógica de negocio",
            "Generar la lista de horarios disponibles para una fecha y trámite específicos",
            "Respetar el horario de atención de la oficina (lun–vie 9:00–17:00, sáb 9:00–13:00, domingo sin atención)",
            "Permitir al Oficial bloquear días o franjas horarias desde el panel",
        ],
        "colaboradores": [
            "→ Cita (consulta citas existentes para detectar empalmes)",
            "→ Tramite (necesita la duración para calcular disponibilidad)",
            "→ HorarioBloqueado (días o franjas bloqueadas)",
        ],
        "nota": (
            "Implementada en citas/servicios/calendario.py (clase Calendario). Usa utilidades de "
            "citas/utils.py para slots y empalmes. La gestión de bloqueos (gestionar_horarios) también "
            "delega en Calendario. El bloqueo de empalmes es lógico; no usa bloqueo pesimista en BD."
        ),
    },
    {
        "numero": 24,
        "titulo": "Generación y validación de QR",
        "nombre": "QR",
        "tipo": "Generación y validación de QR — citas",
        "capa": "Capa: Lógica de Negocio",
        "responsabilidades": [
            "Generar el código QR a partir de la URL de validación de la cita",
            "Guardar la imagen del QR en el campo ImageField del modelo Cita",
            "Validar el token escaneado por el Capturista al llegar el ciudadano",
            "Verificar que la cita esté en estado PENDIENTE antes de marcarla ASISTIDA",
            "Rechazar códigos inexistentes, alterados o ya utilizados con mensaje de error claro",
        ],
        "colaboradores": [
            "→ Cita (guarda token e imagen QR)",
        ],
        "nota": (
            "Implementada en citas/servicios/qr.py (clase QR). La generación de imagen ocurre al guardar "
            "el modelo Cita; la validación y marcado ASISTIDA se hace desde QR.validar_y_marcar_asistida(). "
            "El QR codifica una URL del tipo /citas/validar/{id}/{token}/."
        ),
    },
    {
        "numero": 25,
        "titulo": "Trámite",
        "nombre": "Gestión del catálogo",
        "tipo": "Lógica de negocio — citas",
        "capa": "Capa: Lógica de Negocio",
        "responsabilidades": [
            "Listar trámites activos con costo actualizado para el portal ciudadano",
            "Permitir al Administrador crear, editar, activar y desactivar trámites",
            "Verificar que un trámite no tenga citas futuras en estado PENDIENTE antes de desactivarlo",
            "Registrar en bitácora cada cambio de costo con fecha, valor anterior, valor nuevo y usuario",
            "Proveer el costo al portal cuando el ciudadano cambia de trámite sin recargar la página",
        ],
        "colaboradores": [
            "→ Tramite (gestiona sus registros)",
            "→ SeccionTramite (agrupa los trámites por categoría)",
            "→ Cita (consulta citas futuras)",
            "→ Bitacora (registra cambios en el catálogo)",
        ],
        "nota": (
            "Implementada en citas/servicios/tramite.py (clase Tramite, importada como TramiteNegocio). "
            "El catálogo del admin, la API del portal y las reglas de desactivación delegan en esta clase. "
            "Las señales Django siguen auditando cambios de costo desde el admin."
        ),
    },
    {
        "numero": 26,
        "titulo": "Caja",
        "nombre": "Control de pagos",
        "tipo": "Lógica de negocio — citas",
        "capa": "Capa: Lógica de Negocio",
        "responsabilidades": [
            "Registrar el pago de una cita al momento de su atención en ventanilla",
            "Cambiar el estado de la cita de ASISTIDA a PAGADA al registrar el pago",
            "Permitir que solo Oficial o Administrador cancelen una cita ya pagada",
            "Marcar corte_cierre_listo = True al incluir el pago en el corte del día",
            "Generar el corte de caja al cierre del día desglosado por tipo de trámite",
            "Proveer al panel el monto recaudado en el mes en curso e historial de meses anteriores",
            "Excluir de ingresos los pagos asociados a citas canceladas",
        ],
        "colaboradores": [
            "→ PagoCaja (crea y consulta registros de pago)",
            "→ Cita (obtiene el monto del trámite)",
            "→ UsuarioRegistroCivil (verifica rol)",
            "→ CorteCajaDiario (consolida el cierre diario)",
            "→ Bitacora (registra pagos, cancelaciones y cierres)",
        ],
        "nota": (
            "Implementada en citas/servicios/caja.py (clase Caja). Cobros, reporte financiero y corte "
            "diario delegan en esta clase. La cancelación de pago se modela como cancelación de la cita "
            "pagada; el registro PagoCaja permanece pero deja de contarse en ingresos."
        ),
    },
    {
        "numero": 27,
        "titulo": "Bitácora",
        "nombre": "Auditoría — roles",
        "tipo": "Lógica de negocio — citas",
        "capa": "Capa: Lógica de Negocio",
        "responsabilidades": [
            "Crear registros inmutables en BitacoraAuditoria por cada acción crítica del personal autorizado",
            "Registrar acción, descripción, usuario responsable, fecha, hora e IP",
            "Nunca modificar ni eliminar registros existentes",
            "Permitir consulta con filtros por usuario, tipo de acción y búsqueda libre (descripción, IP, usuario)",
        ],
        "colaboradores": [
            "→ BitacoraAuditoria (persiste los registros)",
        ],
        "nota": (
            "Implementada en citas/servicios/bitacora.py (clase Bitacora). citas/auditoria.py reexporta "
            "funciones que delegan aquí para compatibilidad con señales y middleware. La consulta con "
            "filtros se expone mediante Bitacora.consultar()."
        ),
    },
    {
        "numero": 28,
        "titulo": "Horario bloqueado",
        "nombre": "HorarioBloqueado",
        "tipo": "Modelo — citas",
        "capa": "Capa: Datos",
        "responsabilidades": [
            "Almacenar fecha bloqueada y, opcionalmente, hora específica",
            "Permitir bloqueo de día completo o de una franja horaria",
            "Registrar motivo del bloqueo y fecha de creación",
        ],
        "colaboradores": [
            "→ Calendario / utilidades de horario",
            "→ UsuarioRegistroCivil (creado_por)",
        ],
        "nota": (
            "Permite al Oficial restringir días o horarios no disponibles para agendamiento."
        ),
    },
    {
        "numero": 29,
        "titulo": "Corte caja diario",
        "nombre": "CorteCajaDiario",
        "tipo": "Modelo — citas",
        "capa": "Capa: Datos",
        "responsabilidades": [
            "Almacenar fecha del corte, total recaudado y desglose por trámite",
            "Registrar cantidad de pagos incluidos y usuario que cerró el corte",
            "Ser inmutable una vez creado",
        ],
        "colaboradores": [
            "→ PagoCaja (origen de los montos consolidados)",
            "→ UsuarioRegistroCivil (cerrado_por)",
            "→ BitacoraAuditoria (audita el cierre)",
        ],
        "nota": (
            "Representa el cierre diario de caja con desglose por trámite. Una vez generado, "
            "no puede modificarse ni eliminarse."
        ),
    },
]


def set_cell_shading(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_tarjeta(doc: Document, tarjeta: dict) -> None:
    doc.add_heading(f"Tabla {tarjeta['numero']}", level=1)

    p = doc.add_paragraph()
    run = p.add_run(tarjeta["titulo"])
    run.bold = True
    run.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.add_run(tarjeta["nombre"]).bold = True
    meta.add_run(f"\n{tarjeta['tipo']}\n{tarjeta['capa']}")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(4.0)
    table.columns[1].width = Inches(2.3)

    hdr = table.rows[0].cells
    hdr[0].text = "RESPONSABILIDADES"
    hdr[1].text = "COLABORADORES"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_shading(cell, "E4F1EA")

    max_rows = max(len(tarjeta["responsabilidades"]), len(tarjeta["colaboradores"]))
    for i in range(max_rows):
        row = table.add_row().cells
        resp = tarjeta["responsabilidades"][i] if i < len(tarjeta["responsabilidades"]) else ""
        col = tarjeta["colaboradores"][i] if i < len(tarjeta["colaboradores"]) else ""
        row[0].text = f"• {resp}" if resp else ""
        row[1].text = col

    nota = doc.add_paragraph()
    nota.add_run("Nota: ").bold = True
    nota.add_run(tarjeta["nota"])

    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Diseño de las tarjetas CRC", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "Documento alineado con la implementación actual del sistema Registro Civil 03 "
        "(Django). Incluye la capa de servicios orientada a objetos en citas/servicios/ y "
        "autenticacion/servicios/, con rutas de código reales para cada tarjeta CRC de lógica "
        "de negocio (tablas 16 y 22–27)."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    for tarjeta in TARJETAS:
        add_tarjeta(doc, tarjeta)

    return doc


def main() -> None:
    out_dir = Path(__file__).resolve().parent
    out_path = out_dir / "Diseño de las tarjetas CRC - corregido.docx"
    build_document().save(out_path)
    print(f"Generado: {out_path}")


if __name__ == "__main__":
    main()
